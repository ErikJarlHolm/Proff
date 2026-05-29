"""
WordExport – Word-rapporteksport for ProffAgenten
===================================================
Eksporterer bedriftsvurderinger til Word-filer (.docx) basert på
malen 'Template Bedriftsinformasjon.docx'.

Eksporterer:
    - WORD_TOOL_DEFINITION: JSON-definisjon for Foundry function-calling
    - write_report_to_word(): Hovedfunksjonen som fyller ut malen og lagrer filen

Bruk fra agent.py:
    from .word_export import WORD_TOOL_DEFINITION, write_report_to_word

Avhengigheter:
    - python-docx: Lesing/skriving av .docx-filer
"""

import re
import logging
from datetime import date
from pathlib import Path

from docx import Document

log = logging.getLogger(__name__)

# ── Stier ────────────────────────────────────────────────────────────────────
TEMPLATE_FILE = Path(
    r"C:\Users\erikholm\OneDrive - Atea\Documents\Kunder\Atea AI Norge"
    r"\Technocamp 2026\Template Bedriftsinformasjon.docx"
)
# Fallback: lokal kopi ved siden av koden (OneDrive-filer kan være cloud-only)
TEMPLATE_FILE_LOCAL = Path(__file__).parent.parent / "template.docx"

REPORTS_DIR = Path(
    r"C:\Users\erikholm\OneDrive - Atea\Documents\Kunder\Atea AI Norge"
    r"\Technocamp 2026"
)


def _get_template() -> Path | None:
    """Finn malen – prøv OneDrive først, deretter lokal kopi."""
    for candidate in [TEMPLATE_FILE, TEMPLATE_FILE_LOCAL]:
        if candidate.exists():
            try:
                # Sjekk at filen faktisk er tilgjengelig (ikke cloud-only stub)
                with open(candidate, "rb") as f:
                    f.read(4)
                return candidate
            except (OSError, IOError):
                continue
    return None


# ── Verktøydefinisjon for Foundry ────────────────────────────────────────────

WORD_TOOL_DEFINITION = {
    "type": "function",
    "name": "write_report_to_word",
    "description": (
        "Skriv den ferdige bedriftsvurderingen til en Word-fil basert på mal. "
        "Kall dette verktøyet AUTOMATISK etter at du har samlet inn all informasjon "
        "og formulert vurderingen. Bruk KUN de definerte parameterne nedenfor – "
        "ikke legg til ekstra felter."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "company_name": {
                "type": "string",
                "description": "Bedriftens navn (brukes i overskrift og filnavn).",
            },
            "summary": {
                "type": "string",
                "description": "Kort oppsummering (1–3 setninger). Seksjon 1.",
            },
            "key_facts": {
                "type": "string",
                "description": (
                    "Nøkkelfakta: firmanavn, org.nr, adresse, bransje, ledelse, "
                    "eierskap, antall ansatte, stiftelsesdato, hjemmeside. Seksjon 2."
                ),
            },
            "financials": {
                "type": "string",
                "description": (
                    "Finansiell status: siste regnskap, likviditet, lønnsomhet, "
                    "nøkkeltall, kapital. Seksjon 3."
                ),
            },
            "operations": {
                "type": "string",
                "description": (
                    "Operasjonell vurdering: produkter/tjenester, leveransekjede, "
                    "kapasitet, teknologi. Seksjon 4."
                ),
            },
            "risk_compliance": {
                "type": "string",
                "description": (
                    "Risiko og samsvar: kreditt, rettstvister, regulatoriske forhold, "
                    "konkurs-status, ESG. Seksjon 5."
                ),
            },
            "reputation_market": {
                "type": "string",
                "description": (
                    "Reputasjon og marked: mediedekning, kunder, partnerskap, "
                    "markedsposisjon. Seksjon 6."
                ),
            },
            "recommendations": {
                "type": "string",
                "description": (
                    "Anbefalinger og tiltaksforslag: kort/mellom/langsiktige råd, "
                    "samlet vurdering (grønt/gult/rødt). Seksjon 7."
                ),
            },
            "sources": {
                "type": "string",
                "description": (
                    "Kilder og usikkerhetsvurdering: liste over datakilder med URL/kildetype "
                    "og dato, hva som mangler, begrensninger ved analysen. Seksjon 8."
                ),
            },
        },
        "required": [
            "company_name", "summary", "key_facts", "financials",
            "operations", "risk_compliance", "reputation_market",
            "recommendations", "sources",
        ],
    },
}


# ── Hjelpefunksjoner ─────────────────────────────────────────────────────────


def _slug(text: str) -> str:
    """Lag et filnavnvennlig stikkord fra bedriftsnavnet."""
    text = text.strip().lower()
    text = re.sub(r"[^a-zæøåa-z0-9\s]", "", text)
    text = re.sub(r"\s+", "_", text)
    return text[:40] or "bedrift"


def _insert_text_after_heading(doc: Document, heading_text: str, content: str) -> bool:
    """
    Finn et avsnitt med heading_text og sett inn innhold etter det.

    Støtter både eksakt match og partial match (for tilfeller der malens
    overskrift har ekstra mellomrom eller nummerering).
    """
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        p_text = p.text.strip()
        if p_text == heading_text.strip() or p_text.startswith(heading_text.strip()):
            # Sjekk om neste paragraf er et tomt avsnitt vi kan bruke
            if i + 1 < len(paragraphs) and not paragraphs[i + 1].text.strip():
                target = paragraphs[i + 1]
            else:
                new_p = doc.add_paragraph()
                p._element.addnext(new_p._element)
                target = new_p

            target.clear()
            lines = content.split("\n")
            for idx, line in enumerate(lines):
                run = target.add_run(line)
                if idx < len(lines) - 1:
                    run.add_break()
            return True
    return False


def _set_table_cell_text(doc: Document, table_idx: int, row: int, col: int,
                         prefix: str, value: str) -> None:
    """Erstatt tekst i en tabellcelle, behold format fra malen."""
    try:
        cell = doc.tables[table_idx].rows[row].cells[col]
        for p in cell.paragraphs:
            if prefix in p.text:
                existing_size = None
                for r in p.runs:
                    if r.font.size:
                        existing_size = r.font.size
                        break
                p.clear()
                run = p.add_run(f"{prefix}{value}")
                if existing_size:
                    run.font.size = existing_size
                return
    except (IndexError, AttributeError) as e:
        log.warning("Kunne ikke sette tabellcelle [%d,%d,%d]: %s", table_idx, row, col, e)


# ── Hovedfunksjon ────────────────────────────────────────────────────────────


def write_report_to_word(
    company_name: str,
    summary: str,
    key_facts: str,
    financials: str,
    operations: str,
    risk_compliance: str,
    reputation_market: str,
    recommendations: str,
    sources: str,
) -> dict:
    """
    Fyll ut Word-malen med bedriftsvurdering og lagre som ny fil.

    Returnerer:
        {"status": "ok", "file": "<full filsti>"} ved suksess
        {"error": "<feilmelding>"} ved feil
    """
    template = _get_template()
    if not template:
        return {"error": f"Mal ikke funnet. Forventet: {TEMPLATE_FILE} eller {TEMPLATE_FILE_LOCAL}"}

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    today = str(date.today())
    doc = Document(str(template))

    # ── Tabellfelter ────────────────────────────────────────────────────────
    _set_table_cell_text(doc, 0, 0, 0, "Oppsummering av Bedrift:", f" {company_name}")
    _set_table_cell_text(doc, 1, 0, 0, "Dato:", f" {today}")

    # AI oppsummering (tabell 3)
    try:
        cell = doc.tables[3].rows[0].cells[0]
        for p in cell.paragraphs:
            if "AI oppsummering" in p.text:
                p.clear()
                run = p.add_run(f"AI oppsummering: {summary}")
                break
    except (IndexError, AttributeError):
        pass

    # ── Sett bedriftsnavn i overskrift ───────────────────────────────────────
    for p in doc.paragraphs:
        if p.text.strip() == "Bedrift":
            p.clear()
            run = p.add_run(company_name)
            break

    # ── Innholdsavsnitt ──────────────────────────────────────────────────────
    sections = [
        ("1. Kort oppsummering", summary),
        ("2. Nøkkelfakta", key_facts),
        ("3. Finans", financials),
        ("4. Operasjonell vurdering", operations),
        ("5. Risiko og samsvar", risk_compliance),
        ("6. Reputasjon og marked", reputation_market),
        ("7. Anbefalinger og tiltaksforslag", recommendations),
        ("8. Kilder og usikkerhetsvurdering", sources),
    ]
    for heading, content in sections:
        if content:
            found = _insert_text_after_heading(doc, heading, content)
            if not found:
                log.warning("Overskrift ikke funnet i mal: '%s'", heading)

    # ── Filnavn og lagring ───────────────────────────────────────────────────
    slug = _slug(company_name)
    filename = f"{today}_bedriftsvurdering_{slug}.docx"
    output_path = REPORTS_DIR / filename

    # Unngå overskriving
    counter = 1
    while output_path.exists():
        output_path = REPORTS_DIR / f"{today}_bedriftsvurdering_{slug}_{counter}.docx"
        counter += 1

    doc.save(str(output_path))
    log.info("Rapport lagret: %s", output_path)
    return {"status": "ok", "file": str(output_path)}
